#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Script để tạo file Word chứa Use Case Specifications dưới dạng bảng
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def shade_cell(cell, color):
    """Tô màu cell trong bảng"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._element.get_or_add_tcPr().append(shading_elm)

def create_use_case_document():
    """Tạo document Word với các use case"""
    doc = Document()
    
    # Tiêu đề
    title = doc.add_heading('ĐẶC TẢ USE CASE', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_heading('Hệ Thống Quản Lý Nhà Hàng KFC', level=2)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('Ngày: 19/01/2026 | Phiên Bản: 1.0')
    doc.add_paragraph()
    
    # ========== PHẦN 1: DANH SÁCH USE CASE ==========
    doc.add_heading('1. DANH SÁCH USE CASE TÓNG HỢP', level=1)
    
    # Bảng tóm tắt use case
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Light Grid Accent 1'
    
    # Header
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'ID'
    hdr_cells[1].text = 'Tên Use Case'
    hdr_cells[2].text = 'Actor Chính'
    hdr_cells[3].text = 'Module'
    hdr_cells[4].text = 'Mô Tả Ngắn'
    
    # Tô màu header
    for cell in hdr_cells:
        shade_cell(cell, 'D3D3D3')
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    # Dữ liệu use case
    use_cases = [
        ('UC01', 'Cất hàng trực tuyến', 'Khách hàng', 'Đặt hàng & Thanh toán', 'Khách hàng chọn sản phẩm, thêm vào giỏ, đặt hàng'),
        ('UC02', 'Thanh toán online', 'Khách hàng', 'Đặt hàng & Thanh toán', 'Thanh toán qua cổng thanh toán online'),
        ('UC03', 'Gửi đơn xuống bếp', 'Nhân viên', 'Đặt hàng & Thanh toán', 'Gửi đơn cho bếp chế biến sau khi thanh toán'),
        ('UC04', 'Thanh toán tiền mặt', 'Khách hàng', 'Đặt hàng & Thanh toán', 'Thanh toán tiền mặt khi nhận hàng'),
        ('UC05', 'Quản lý đơn hàng', 'Nhân viên', 'Vận hành đơn hàng', 'Xem, lọc, quản lý danh sách đơn hàng'),
        ('UC06', 'Hủy đơn hàng', 'Khách hàng/Nhân viên', 'Vận hành đơn hàng', 'Hủy đơn hàng và hoàn tiền nếu cần'),
        ('UC07', 'Tìm kiếm đơn hàng', 'Khách hàng/Nhân viên', 'Vận hành đơn hàng', 'Tìm kiếm đơn hàng theo mã, ngày, trạng thái'),
        ('UC08', 'Cập nhật trạng thái đơn', 'Nhân viên/Đầu bếp', 'Vận hành đơn hàng', 'Cập nhật trạng thái: Đang chuẩn bị → Giao → Hoàn thành'),
        ('UC09', 'Quản lý sản phẩm', 'Quản lý', 'Quản lý sản phẩm', 'Thêm, sửa, xóa sản phẩm'),
        ('UC10', 'Quản lý danh mục', 'Quản lý', 'Quản lý sản phẩm', 'Quản lý các danh mục sản phẩm'),
        ('UC11', 'Quản lý thực đơn', 'Quản lý', 'Quản lý sản phẩm', 'Thêm/xóa sản phẩm vào thực đơn'),
        ('UC12', 'Đăng nhập', 'Tất cả người dùng', 'Quản lý người dùng', 'Đăng nhập vào hệ thống'),
        ('UC13', 'Đăng ký', 'Khách hàng', 'Quản lý người dùng', 'Tạo tài khoản mới'),
        ('UC14', 'Phân quyền người dùng', 'Quản lý', 'Quản lý người dùng', 'Gán vai trò và quyền cho người dùng'),
        ('UC15', 'Xem báo cáo', 'Quản lý', 'Báo cáo & Thống kê', 'Xem báo cáo doanh thu, sản phẩm, đơn hàng'),
        ('UC16', 'Quản lý tài chính', 'Quản lý tài chính', 'Báo cáo & Thống kê', 'Quản lý doanh thu, giao dịch'),
        ('UC17', 'Xác thực người dùng', 'Hệ thống', 'Hệ thống chung', 'Xác thực token JWT'),
    ]
    
    for uc_id, name, actor, module, desc in use_cases:
        row_cells = table.add_row().cells
        row_cells[0].text = uc_id
        row_cells[1].text = name
        row_cells[2].text = actor
        row_cells[3].text = module
        row_cells[4].text = desc
    
    doc.add_paragraph()
    
    # ========== PHẦN 2: CHI TIẾT USE CASE ==========
    doc.add_heading('2. CHI TIẾT TỪNG USE CASE', level=1)
    
    # Use case details
    details = [
        {
            'id': 'UC01',
            'name': 'Cất Hàng Trực Tuyến',
            'description': 'Cung cấp giao diện cho khách hàng lựa chọn các món ăn, combo từ thực đơn của KFC và thêm vào giỏ hàng để tiến hành đặt hàng trực tuyến.',
            'actor': 'Khách hàng',
            'precondition': 'Khách hàng đã đăng nhập, có sản phẩm trong giỏ hàng',
            'postcondition': 'Đơn hàng được tạo, đợi thanh toán',
            'steps': [
                '1. Khách hàng chọn sản phẩm từ menu',
                '2. Thêm sản phẩm vào giỏ hàng',
                '3. Xem giỏ hàng và kiểm tra thông tin',
                '4. Nhập/xác nhận thông tin giao hàng',
                '5. Chọn phương thức thanh toán',
                '6. Xác nhận đặt hàng',
                '7. Hệ thống tạo đơn hàng với trạng thái "Chờ thanh toán"'
            ],
            'exceptions': [
                'E1: Sản phẩm hết hàng → Thông báo và gợi ý sản phẩm khác',
                'E2: Thông tin giao hàng không hợp lệ → Yêu cầu nhập lại',
                'E3: Giỏ hàng trống → Không cho phép đặt hàng'
            ]
        },
        {
            'id': 'UC02',
            'name': 'Thanh Toán Online',
            'description': 'Cho phép khách hàng thực hiện thanh toán đơn hàng thông qua các cổng thanh toán điện tử tích hợp để hoàn tất quá trình mua hàng.',
            'actor': 'Khách hàng',
            'precondition': 'Đơn hàng tồn tại, trạng thái "Chờ thanh toán"',
            'postcondition': 'Thanh toán thành công, đơn hàng chuyển trạng thái "Đã xác nhận"',
            'steps': [
                '1. Khách hàng chọn phương thức thanh toán online',
                '2. Điều hướng đến cổng thanh toán',
                '3. Nhập thông tin thẻ/tài khoản',
                '4. Xác nhận thanh toán',
                '5. Cổng thanh toán phản hồi kết quả',
                '6. Hệ thống cập nhật trạng thái đơn hàng'
            ],
            'exceptions': [
                'E1: Thẻ bị từ chối → Yêu cầu thử lại hoặc đổi phương thức',
                'E2: Timeout kết nối → Lưu lại trang thanh toán, cho phép thử lại',
                'E3: Lỗi cổng → Thông báo lỗi, hướng dẫn liên hệ hỗ trợ'
            ]
        },
        {
            'id': 'UC03',
            'name': 'Gửi Đơn Xuống Bếp',
            'description': 'Hệ thống tự động hoặc nhân viên xác nhận chuyển thông tin đơn hàng đã thanh toán đến bộ phận bếp để bắt đầu quy trình chế biến món ăn.',
            'actor': 'Nhân viên',
            'precondition': 'Đơn hàng được thanh toán thành công',
            'postcondition': 'Đơn hàng được gửi đến bếp, trạng thái "Đang chuẩn bị"',
            'steps': [
                '1. Hệ thống nhận đơn hàng đã thanh toán',
                '2. Tạo ticket in/gửi đơn xuống bếp',
                '3. Gửi đơn đến bếp',
                '4. Đầu bếp nhận đơn',
                '5. Cập nhật trạng thái đơn hàng trong hệ thống'
            ],
            'exceptions': [
                'E1: Nguyên liệu không đủ → Nhân viên thông báo, hủy/thay thế sản phẩm',
                'E2: Máy in hỏng → Hệ thống ghi log, gửi alert'
            ]
        },
        {
            'id': 'UC05',
            'name': 'Quản Lý Đơn Hàng',
            'description': 'Cho phép nhân viên quản lý theo dõi, lọc và cập nhật thông tin các đơn hàng trong hệ thống để đảm bảo quy trình phục vụ diễn ra suôn sẻ.',
            'actor': 'Nhân viên',
            'precondition': 'Nhân viên đã đăng nhập',
            'postcondition': 'Đơn hàng được cập nhật trạng thái',
            'steps': [
                '1. Nhân viên truy cập trang quản lý đơn hàng',
                '2. Hệ thống hiển thị danh sách đơn hàng',
                '3. Nhân viên lọc/tìm kiếm theo tiêu chí',
                '4. Chọn đơn hàng và xem chi tiết',
                '5. Cập nhật thông tin đơn hàng (chi tiết xem UC08)',
                '6. Lưu thay đổi'
            ],
            'exceptions': [
                'E1: Đơn hàng không tìm thấy → Thông báo lỗi',
                'E2: Quyền hạn không đủ → Từ chối truy cập'
            ]
        },
        {
            'id': 'UC09',
            'name': 'Quản Lý Sản Phẩm',
            'description': 'Cung cấp công cụ cho Quản lý để thực hiện các thao tác thêm mới, chỉnh sửa hoặc xóa bỏ các thông tin về món ăn, combo và giá cả trên hệ thống.',
            'actor': 'Quản lý',
            'precondition': 'Quản lý đã đăng nhập vào phần admin',
            'postcondition': 'Sản phẩm được thêm/sửa/xóa thành công',
            'steps': [
                '1. Quản lý truy cập trang quản lý sản phẩm',
                '2. Hệ thống hiển thị danh sách sản phẩm',
                '3. Quản lý chọn thao tác: Thêm / Sửa / Xóa sản phẩm',
                '4. Nhập hoặc cập nhật thông tin sản phẩm (Tên, Mô tả, Giá, Hình ảnh, Danh mục, v.v.)',
                '5. Xác nhận thao tác',
                '6. Hệ thống cập nhật CSDL'
            ],
            'exceptions': [
                'E1: Thông tin không đầy đủ → Yêu cầu nhập lại',
                'E2: Hình ảnh quá lớn → Yêu cầu nén hoặc chọn hình khác'
            ]
        },
        {
            'id': 'UC12',
            'name': 'Đăng Nhập',
            'description': 'Xác thực danh tính người dùng vào hệ thống dựa trên thông tin tài khoản đã đăng ký để cấp quyền truy cập tương ứng với vai trò.',
            'actor': 'Tất cả người dùng',
            'precondition': 'Người dùng có tài khoản',
            'postcondition': 'Người dùng được cấp quyền truy cập',
            'steps': [
                '1. Người dùng truy cập trang đăng nhập',
                '2. Nhập email/username',
                '3. Nhập mật khẩu',
                '4. Nhấn "Đăng nhập"',
                '5. Hệ thống xác thực thông tin',
                '6. Tạo session/token cho người dùng',
                '7. Điều hướng đến trang chính'
            ],
            'exceptions': [
                'E1: Email/username không tồn tại → Thông báo lỗi',
                'E2: Mật khẩu sai → Thông báo lỗi, đếm lần sai',
                'E3: Tài khoản bị khóa → Thông báo và yêu cầu liên hệ hỗ trợ',
                'E4: Quá 5 lần sai mật khẩu → Khóa tài khoản tạm thời'
            ]
        },
        {
            'id': 'UC14',
            'name': 'Phân Quyền Người Dùng',
            'description': 'Cho phép Quản lý thiết lập và điều chỉnh các quyền hạn, vai trò cho từng thành viên trong hệ thống để bảo mật và tối ưu quy trình làm việc.',
            'actor': 'Quản lý',
            'precondition': 'Quản lý đã đăng nhập, có danh sách người dùng',
            'postcondition': 'Quyền của người dùng được cập nhật thành công',
            'steps': [
                '1. Quản lý truy cập trang quản lý người dùng',
                '2. Hệ thống hiển thị danh sách tất cả người dùng',
                '3. Quản lý chọn một người dùng',
                '4. Xem các vai trò/quyền hiện tại của người dùng',
                '5. Chọn vai trò cần gán (Khách hàng, Nhân viên, Thủ kho, Đầu bếp, Quản lý)',
                '6. Chọn các quyền cụ thể cho vai trò đó',
                '7. Xác nhận thay đổi',
                '8. Hệ thống cập nhật CSDL và ghi log'
            ],
            'exceptions': [
                'E1: Người dùng không tìm thấy → Thông báo lỗi',
                'E2: Quyền hạn không đủ → Từ chối thao tác',
                'E3: Cố gắng xóa quyền admin duy nhất → Thông báo cảnh báo'
            ]
        },
        {
            'id': 'UC15',
            'name': 'Xem Báo Cáo',
            'description': 'Tổng hợp dữ liệu từ hệ thống để tạo ra các báo cáo trực quan về doanh thu, sản phẩm và hiệu suất bán hàng cho Quản lý.',
            'actor': 'Quản lý',
            'precondition': 'Quản lý đã đăng nhập',
            'postcondition': 'Báo cáo được hiển thị cho quản lý',
            'steps': [
                '1. Quản lý truy cập trang báo cáo',
                '2. Chọn loại báo cáo cần xem (Doanh thu, Sản phẩm, Đơn hàng, Người dùng)',
                '3. Chọn khoảng thời gian (Ngày, Tuần, Tháng, Năm)',
                '4. Có thể lọc theo điều kiện khác (Sản phẩm, Danh mục, Nhân viên)',
                '5. Hệ thống tính toán và hiển thị dữ liệu',
                '6. Quản lý có thể xem biểu đồ, bảng thống kê chi tiết',
                '7. Có tùy chọn xuất báo cáo dưới dạng PDF/Excel'
            ],
            'exceptions': [
                'E1: Không có dữ liệu trong khoảng thời gian → Hiển thị "Không có dữ liệu"',
                'E2: Lỗi truy vấn CSDL → Thông báo lỗi, yêu cầu thử lại',
                'E3: Quyền hạn không đủ → Chỉ xem báo cáo của bộ phận của mình'
            ]
        },
    ]
    
    for detail in details:
        # Tiêu đề use case
        doc.add_heading(f"{detail['id']}: {detail['name']}", level=2)
        
        # Bảng thông tin
        info_table = doc.add_table(rows=1, cols=2)
        info_table.style = 'Light Grid Accent 1'
        
        rows = [
            ('ID', detail['id']),
            ('Tên Use Case', detail['name']),
            ('Mô Tả', detail['description']),
            ('Actor Chính', detail['actor']),
            ('Tiền Điều Kiện', detail['precondition']),
            ('Hậu Điều Kiện', detail['postcondition']),
        ]
        
        for label, value in rows:
            row_cells = info_table.add_row().cells
            row_cells[0].text = label
            row_cells[1].text = value
            shade_cell(row_cells[0], 'E8E8E8')
            row_cells[0].paragraphs[0].runs[0].font.bold = True
        
        # Luồng chính
        flow_steps = '\n'.join(detail['steps'])
        row_cells = info_table.add_row().cells
        row_cells[0].text = 'Luồng Chính'
        row_cells[1].text = flow_steps
        shade_cell(row_cells[0], 'E8E8E8')
        row_cells[0].paragraphs[0].runs[0].font.bold = True
        
        # Ngoại lệ
        if detail['exceptions']:
            exceptions_text = '\n'.join(detail['exceptions'])
            row_cells = info_table.add_row().cells
            row_cells[0].text = 'Luồng Ngoại Lệ'
            row_cells[1].text = exceptions_text
            shade_cell(row_cells[0], 'E8E8E8')
            row_cells[0].paragraphs[0].runs[0].font.bold = True
        
        doc.add_paragraph()
    
    # ========== PHẦN 3: ACTORS ==========
    doc.add_page_break()
    doc.add_heading('3. DANH SÁCH ACTORS', level=1)
    
    actors_table = doc.add_table(rows=1, cols=2)
    actors_table.style = 'Light Grid Accent 1'
    
    hdr_cells = actors_table.rows[0].cells
    hdr_cells[0].text = 'Actor'
    hdr_cells[1].text = 'Mô Tả'
    
    for cell in hdr_cells:
        shade_cell(cell, 'D3D3D3')
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    actors = [
        ('Khách Hàng', 'Người dùng cuối mua hàng qua website'),
        ('Nhân Viên', 'Nhân viên xử lý, hỗ trợ khách hàng'),
        ('Thủ Kho', 'Quản lý tồn kho, nguyên liệu'),
        ('Đầu Bếp', 'Người nấu ăn, xử lý đơn hàng'),
        ('Quản Lý', 'Người quản lý hệ thống, báo cáo'),
    ]
    
    for actor, desc in actors:
        row_cells = actors_table.add_row().cells
        row_cells[0].text = actor
        row_cells[1].text = desc
    
    # ========== PHẦN 4: TRẠNG THÁI ĐƠN HÀNG ==========
    doc.add_page_break()
    doc.add_heading('4. TRẠNG THÁI ĐƠN HÀNG', level=1)
    
    status_table = doc.add_table(rows=1, cols=3)
    status_table.style = 'Light Grid Accent 1'
    
    hdr_cells = status_table.rows[0].cells
    hdr_cells[0].text = 'Trạng Thái'
    hdr_cells[1].text = 'Mô Tả'
    hdr_cells[2].text = 'Trạng Thái Tiếp Theo'
    
    for cell in hdr_cells:
        shade_cell(cell, 'D3D3D3')
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    statuses = [
        ('Chờ Thanh Toán', 'Đơn hàng vừa được tạo, chờ thanh toán', 'Đã xác nhận / Đã hủy'),
        ('Đã Xác Nhận', 'Thanh toán thành công, chờ xử lý', 'Đang chuẩn bị / Đã hủy'),
        ('Đang Chuẩn Bị', 'Đầu bếp đang chuẩn bị đơn hàng', 'Đang giao / Đã hủy'),
        ('Đang Giao', 'Đơn hàng đang được giao cho khách', 'Hoàn thành / Giao thất bại'),
        ('Hoàn Thành', 'Khách hàng đã nhận đơn hàng', 'Không (trạng thái cuối)'),
        ('Đã Hủy', 'Đơn hàng được hủy', 'Không (trạng thái cuối)'),
    ]
    
    for status, desc, next_status in statuses:
        row_cells = status_table.add_row().cells
        row_cells[0].text = status
        row_cells[1].text = desc
        row_cells[2].text = next_status
    
    # ========== PHẦN 5: YÊU CẦU PHI CHỨC NĂNG ==========
    doc.add_page_break()
    doc.add_heading('5. YÊU CẦU PHI CHỨC NĂNG', level=1)
    
    nfr_table = doc.add_table(rows=1, cols=2)
    nfr_table.style = 'Light Grid Accent 1'
    
    hdr_cells = nfr_table.rows[0].cells
    hdr_cells[0].text = 'Yêu Cầu'
    hdr_cells[1].text = 'Chi Tiết'
    
    for cell in hdr_cells:
        shade_cell(cell, 'D3D3D3')
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    nfrs = [
        ('Hiệu Năng', 'Thời gian phản hồi < 3 giây cho 90% request'),
        ('Sẵn Sàng', 'Hệ thống hoạt động 24/7 với 99% uptime'),
        ('Bảo Mật', 'Mã hóa HTTPS, Hash mật khẩu, Token JWT'),
        ('Khả Năng Mở Rộng', 'Hỗ trợ ít nhất 10.000 users đồng thời'),
        ('Khả Năng Bảo Trì', 'Log chi tiết, giám sát hệ thống'),
        ('Tương Thích', 'Hoạt động trên Chrome, Firefox, Safari, Edge'),
    ]
    
    for nfr, detail in nfrs:
        row_cells = nfr_table.add_row().cells
        row_cells[0].text = nfr
        row_cells[1].text = detail
    
    # Footer
    doc.add_page_break()
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_para.add_run('---').font.size = Pt(12)
    
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run('Tài Liệu Được Tạo: 19/01/2026 | Phiên Bản: 1.0 | Trạng Thái: Hoàn thành ban đầu')
    
    # Lưu file
    output_path = 'e:\\Code\\PTIT\\BTL\\CNPM + CSDLPT\\kfc\\docs\\USE_CASE_SPECIFICATIONS.docx'
    doc.save(output_path)
    print(f"✅ File Word đã được tạo thành công: {output_path}")

if __name__ == '__main__':
    create_use_case_document()
